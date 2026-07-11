from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
from pathlib import Path

def image_dimensions(path: Path) -> tuple[int, int]:
    try:
        from PIL import Image
        with Image.open(path) as image:
            return image.size
    except Exception:
        return (1600, 900)


def effective_units(text: str) -> int:
    text = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)
    text = re.sub(r"!\[[^]]*\]\([^)]+\)", " ", text)
    text = re.sub(r"[#*_>`|~-]", " ", text)
    return len(re.findall(r"[\u4e00-\u9fff]", text)) + len(re.findall(r"\b[A-Za-z][A-Za-z'-]*\b", text))


def export(workspace: Path) -> dict:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except Exception as exc:
        raise SystemExit("python-docx is required: pip install python-docx") from exc
    state = json.loads((workspace / "状态" / "工作流状态.json").read_text(encoding="utf-8"))
    source = workspace / "论文" / "论文正文.md"
    output = workspace / "论文" / "数模论文.docx"
    if not source.exists():
        raise SystemExit(f"missing source: {source}")
    text = source.read_text(encoding="utf-8")
    document = Document()
    section = document.sections[0]
    if state.get("competition") == "mcm-icm":
        section.page_width, section.page_height = Cm(21.59), Cm(27.94)
    else:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.54)
    section.top_margin = section.bottom_margin = Cm(2.54)
    usable_width_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    usable_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / 360000
    document.styles["Normal"].font.name = "Times New Roman" if state.get("competition") == "mcm-icm" else "宋体"
    document.styles["Normal"].font.size = Pt(12 if state.get("competition") == "mcm-icm" else 10.5)
    images: list[dict] = []
    for raw in text.splitlines():
        line = raw.strip()
        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", line)
        if image_match:
            image_path = Path(image_match.group(2))
            if not image_path.is_absolute():
                image_path = (source.parent / image_path).resolve()
            if not image_path.exists():
                raise SystemExit(f"missing image: {image_path}")
            px_w, px_h = image_dimensions(image_path)
            ratio = px_h / max(px_w, 1)
            width_cm = min(15.0, float(usable_width_cm))
            height_cm = width_cm * ratio
            max_height_cm = min(20.0, float(usable_height_cm) * 0.72)
            if height_cm > max_height_cm:
                height_cm = max_height_cm
                width_cm = height_cm / max(ratio, 0.01)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
            if image_match.group(1):
                caption = document.add_paragraph(image_match.group(1))
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            images.append({"path": str(image_path), "width_cm": round(width_cm, 3), "height_cm": round(height_cm, 3)})
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            document.add_heading(heading.group(2), level=min(len(heading.group(1)), 3))
        elif line:
            document.add_paragraph(line)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    report = {"output_format": "docx", "competition": state.get("competition"), "source": str(source), "source_sha256": hashlib.sha256(source.read_bytes()).hexdigest(), "source_mtime_ns": source.stat().st_mtime_ns, "output": str(output), "effective_body_units": effective_units(text), "usable_width_cm": round(float(usable_width_cm), 3), "usable_height_cm": round(float(usable_height_cm), 3), "images": images, "page_count": None, "preview_pdf": None}
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        preview_dir = workspace / "论文" / "docx_preview"
        preview_dir.mkdir(parents=True, exist_ok=True)
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(preview_dir), str(output)], capture_output=True, text=True)
        preview = preview_dir / "数模论文.pdf"
        if preview.exists():
            report["preview_pdf"] = str(preview)
            try:
                from pypdf import PdfReader
                report["page_count"] = len(PdfReader(str(preview)).pages)
            except Exception:
                pass
    (workspace / "论文" / "docx_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Export Markdown manuscript to bounded DOCX.")
    parser.add_argument("--workspace", default=".")
    args = parser.parse_args()
    export(Path(args.workspace).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
