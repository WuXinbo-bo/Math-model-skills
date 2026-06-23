#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Revision Toolkit — 通用 DOCX 修订工具链

功能：
1. analyze  — 分析 DOCX 结构（公式、图片、表格、段落）
2. map      — 映射关键段落位置
3. replace  — 精确替换文本（保留 OMML 公式和图片）
4. diff     — 比较两个 DOCX 文件差异
5. batch    — 批量执行修订操作

整合自: edit_inplace.py, compare_docx.py, map_docx.py
"""

import sys
import io
import re
import json
import zipfile
import argparse
from pathlib import Path
from copy import deepcopy

# Ensure UTF-8 output
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════
# Core Functions
# ══════════════════════════════════════════════════════════════

def analyze_docx(path: str) -> dict:
    """
    分析 DOCX 文件结构，返回元数据。
    
    Returns:
        dict: 包含路径、大小、公式数、图片数、表格数、字符数等
    """
    path = Path(path)
    if not path.exists():
        return {"path": str(path), "exists": False}
    
    size = path.stat().st_size
    
    try:
        from docx import Document
        doc = Document(str(path))
        
        # Count paragraphs, tables, images
        paragraphs = len(doc.paragraphs)
        tables = len(doc.tables)
        
        # Count equations and images via XML
        with zipfile.ZipFile(str(path), "r") as z:
            names = z.namelist()
            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
            
        omml_count = len(re.findall(r"<m:oMath", xml))
        images = [n for n in names if "media" in n.lower() or n.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
        drawings = len(re.findall(r"<wp:inline|<wp:anchor", xml))
        raw_dollar = len(re.findall(r"<w:t[^>]*>\$", xml))
        
        # Count characters
        chars = 0
        for para in doc.paragraphs:
            # Skip OMML runs
            for run in para.runs:
                if '<m:oMath' not in run._element.xml and run.text:
                    chars += len(run.text)
        
        # First 200 chars preview
        preview = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 20:
                preview = text[:200]
                break
        
        return {
            "path": path.name,
            "size_kb": round(size / 1024, 1),
            "paragraphs": paragraphs,
            "tables": tables,
            "omml_equations": omml_count,
            "images": len(images),
            "image_names": [Path(n).name for n in images[:5]],
            "drawings": drawings,
            "chars": chars,
            "raw_dollar_signs": raw_dollar,
            "preview": preview,
        }
    except ImportError:
        return {"path": str(path), "exists": True, "error": "python-docx not installed"}
    except Exception as e:
        return {"path": str(path), "exists": True, "error": str(e)}


def map_paragraphs(path: str, keywords: list = None) -> list:
    """
    映射 DOCX 段落结构，找到关键段落位置。
    
    Args:
        path: DOCX 文件路径
        keywords: 要搜索的关键词列表，None 则返回所有段落
    
    Returns:
        list: 段落信息列表 [{index, style, text, preview_prev, preview_next}]
    """
    from docx import Document
    doc = Document(path)
    results = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # If no keywords, return all non-empty paragraphs
        if keywords is None:
            style = para.style.name if para.style else "None"
            prev_text = doc.paragraphs[i-1].text[:50] if i > 0 else "(START)"
            next_text = doc.paragraphs[i+1].text[:50] if i+1 < len(doc.paragraphs) else "(END)"
            results.append({
                "index": i,
                "style": style,
                "text": text[:100],
                "preview_prev": prev_text,
                "preview_next": next_text,
            })
            continue
        
        # Check if any keyword matches
        for keyword in keywords:
            if keyword in text:
                style = para.style.name if para.style else "None"
                prev_text = doc.paragraphs[i-1].text[:50] if i > 0 else "(START)"
                next_text = doc.paragraphs[i+1].text[:50] if i+1 < len(doc.paragraphs) else "(END)"
                results.append({
                    "index": i,
                    "style": style,
                    "text": text[:100],
                    "keyword_match": keyword,
                    "preview_prev": prev_text,
                    "preview_next": next_text,
                })
                break
    
    return results


def get_para_text(para) -> str:
    """获取段落可见文本（跳过 OMML 公式）"""
    texts = []
    for run in para.runs:
        if '<m:oMath' not in run._element.xml and run.text:
            texts.append(run.text)
    return ''.join(texts)


def replace_in_para(para, old_text: str, new_text: str) -> bool:
    """
    在段落中替换文本，仅修改非 OMML 的 run。
    
    Returns:
        bool: 是否成功替换
    """
    full_text = ''.join(r.text for r in para.runs if r.text)
    if old_text not in full_text:
        return False
    
    # Simple approach: find the run containing the old text and replace
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    
    # Fallback: multi-run replacement
    remaining = new_text
    for run in para.runs:
        if run.text and old_text:
            if old_text.startswith(run.text):
                old_text = old_text[len(run.text):]
                run.text = remaining[:len(run.text)]
                remaining = remaining[len(run.text):]
            elif run.text in old_text:
                run.text = run.text.replace(run.text, remaining[:len(run.text)])
                remaining = remaining[len(run.text):]
                old_text = old_text.replace(run.text[:len(run.text)], '', 1)
    return True


def replace_text_in_docx(
    docx_path: str,
    replacements: list,
    output_path: str = None,
    preserve_omml: bool = True
) -> dict:
    """
    批量替换 DOCX 中的文本，保留 OMML 公式和图片。
    
    Args:
        docx_path: 源 DOCX 路径
        replacements: 替换列表 [{"old": "旧文本", "new": "新文本"}]
        output_path: 输出路径，None 则覆盖原文件
        preserve_omml: 是否保留 OMML 公式（默认 True）
    
    Returns:
        dict: {success, changes_applied, output_path}
    """
    from docx import Document
    
    doc = Document(docx_path)
    changes = []
    errors = []
    
    for i, rep in enumerate(replacements):
        old_text = rep.get("old", "")
        new_text = rep.get("new", "")
        
        if not old_text:
            errors.append(f"Replacement {i+1}: empty old_text")
            continue
        
        found = False
        for para in doc.paragraphs:
            para_text = get_para_text(para)
            if old_text in para_text:
                if replace_in_para(para, old_text, new_text):
                    changes.append({
                        "index": i + 1,
                        "old": old_text[:50],
                        "new": new_text[:50],
                        "status": "applied"
                    })
                    found = True
                    break
        
        if not found:
            errors.append(f"Replacement {i+1}: '{old_text[:30]}...' not found")
            changes.append({
                "index": i + 1,
                "old": old_text[:50],
                "new": new_text[:50],
                "status": "not_found"
            })
    
    # Save
    out = output_path or docx_path
    doc.save(out)
    
    # Verify
    with zipfile.ZipFile(out, "r") as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
        omml = len(re.findall(r"<m:oMath", xml))
        raw_dollar = len(re.findall(r"<w:t[^>]*>\$", xml))
    
    return {
        "success": True,
        "output_path": out,
        "changes_applied": len([c for c in changes if c["status"] == "applied"]),
        "changes_failed": len([c for c in changes if c["status"] != "applied"]),
        "omml_preserved": omml,
        "raw_dollar_signs": raw_dollar,
        "details": changes,
        "errors": errors,
    }


def diff_docx(path_a: str, path_b: str) -> dict:
    """
    比较两个 DOCX 文件的文本差异。
    
    Returns:
        dict: {added, removed, modified, stats}
    """
    from docx import Document
    
    def extract_text(path):
        doc = Document(path)
        texts = []
        for para in doc.paragraphs:
            text = get_para_text(para).strip()
            if text:
                texts.append(text)
        return texts
    
    texts_a = extract_text(path_a)
    texts_b = extract_text(path_b)
    
    set_a = set(texts_a)
    set_b = set(texts_b)
    
    added = list(set_b - set_a)[:20]  # Max 20 items
    removed = list(set_a - set_b)[:20]
    
    # Find modified (similar but not identical)
    modified = []
    for t_a in texts_a:
        for t_b in texts_b:
            if t_a != t_b and t_a[:20] == t_b[:20] and len(t_a) > 20:
                modified.append({"old": t_a[:80], "new": t_b[:80]})
                if len(modified) >= 10:
                    break
        if len(modified) >= 10:
            break
    
    return {
        "file_a": Path(path_a).name,
        "file_b": Path(path_b).name,
        "paragraphs_a": len(texts_a),
        "paragraphs_b": len(texts_b),
        "added": added,
        "removed": removed,
        "modified": modified,
        "stats": {
            "added_count": len(added),
            "removed_count": len(removed),
            "modified_count": len(modified),
        }
    }


def batch_replace_from_json(json_path: str, docx_path: str = None, output_path: str = None) -> dict:
    """
    从 JSON 文件批量执行修订操作。
    
    JSON 格式:
    {
        "source": "path/to/original.docx",
        "output": "path/to/output.docx",
        "replacements": [
            {"old": "旧文本1", "new": "新文本1"},
            {"old": "旧文本2", "new": "新文本2"}
        ]
    }
    """
    with open(json_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    
    source = docx_path or config.get("source")
    output = output_path or config.get("output")
    replacements = config.get("replacements", [])
    
    if not source:
        return {"success": False, "error": "No source DOCX specified"}
    
    return replace_text_in_docx(source, replacements, output)


def verify_docx_integrity(path: str) -> dict:
    """
    验证 DOCX 文件完整性（图片、公式是否完好）。
    
    Returns:
        dict: 验证结果
    """
    path = Path(path)
    if not path.exists():
        return {"valid": False, "error": "File not found"}
    
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            # Check ZIP integrity
            bad = z.testzip()
            if bad:
                return {"valid": False, "error": f"Corrupted file in ZIP: {bad}"}
            
            names = z.namelist()
            
            # Check required parts
            required = ["word/document.xml", "[Content_Types].xml"]
            missing = [r for r in required if r not in names]
            if missing:
                return {"valid": False, "error": f"Missing required parts: {missing}"}
            
            # Analyze content
            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
            
            return {
                "valid": True,
                "size_kb": round(path.stat().st_size / 1024, 1),
                "zip_entries": len(names),
                "omml_equations": len(re.findall(r"<m:oMath", xml)),
                "images": len([n for n in names if "media" in n.lower()]),
                "raw_dollar": len(re.findall(r"<w:t[^>]*>\$", xml)),
            }
    except Exception as e:
        return {"valid": False, "error": str(e)}


# ══════════════════════════════════════════════════════════════
# CLI Commands
# ══════════════════════════════════════════════════════════════

def cmd_analyze(args):
    """分析 DOCX 文件结构"""
    import json
    
    if hasattr(args, 'files') and args.files:
        # Multiple files
        results = [analyze_docx(f) for f in args.files]
    else:
        results = [analyze_docx(args.file)]
    
    print(json.dumps(results if len(results) > 1 else results[0], 
                     ensure_ascii=False, indent=2))


def cmd_map(args):
    """映射 DOCX 段落结构"""
    import json
    
    keywords = args.keywords.split(",") if args.keywords else None
    results = map_paragraphs(args.file, keywords)
    
    print(json.dumps(results, ensure_ascii=False, indent=2))


def cmd_replace(args):
    """执行文本替换"""
    import json
    
    if args.json_file:
        # Batch mode from JSON
        result = batch_replace_from_json(args.json_file, args.file, args.output)
    elif args.old and args.new is not None:
        # Single replacement
        result = replace_text_in_docx(
            args.file, 
            [{"old": args.old, "new": args.new}], 
            args.output
        )
    else:
        print("Error: Specify --json-file or --old/--new")
        sys.exit(1)
    
    print(json.dumps(result, ensure_ascii=False, indent=2))
    sys.exit(0 if result.get("success") else 1)


def cmd_diff(args):
    """比较两个 DOCX 差异"""
    import json
    
    result = diff_docx(args.file_a, args.file_b)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def cmd_verify(args):
    """验证 DOCX 完整性"""
    import json
    
    result = verify_docx_integrity(args.file)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    sys.exit(0 if result.get("valid") else 1)


# ══════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="DOCX Revision Toolkit — 通用 DOCX 修订工具链",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 分析 DOCX 结构
  python docx_revision.py analyze paper.docx
  
  # 映射关键段落
  python docx_revision.py map paper.docx --keywords "摘要,模型,结论"
  
  # 单个替换
  python docx_revision.py replace paper.docx --old "旧文本" --new "新文本"
  
  # 批量替换（从 JSON）
  python docx_revision.py replace paper.docx --json-file revisions.json
  
  # 比较两个版本
  python docx_revision.py diff paper_v1.docx paper_v2.docx
  
  # 验证文件完整性
  python docx_revision.py verify paper.docx
        """
    )
    
    sub = parser.add_subparsers(dest="command", required=True)
    
    # analyze
    p_analyze = sub.add_parser("analyze", help="分析 DOCX 文件结构")
    p_analyze.add_argument("file", nargs="?", help="DOCX 文件路径")
    p_analyze.add_argument("--files", nargs="+", help="多个 DOCX 文件（比较模式）")
    
    # map
    p_map = sub.add_parser("map", help="映射 DOCX 段落结构")
    p_map.add_argument("file", help="DOCX 文件路径")
    p_map.add_argument("--keywords", help="搜索关键词，逗号分隔")
    
    # replace
    p_replace = sub.add_parser("replace", help="执行文本替换")
    p_replace.add_argument("file", help="DOCX 文件路径")
    p_replace.add_argument("--old", help="要替换的旧文本")
    p_replace.add_argument("--new", help="新文本")
    p_replace.add_argument("--json-file", help="批量替换 JSON 文件")
    p_replace.add_argument("--output", "-o", help="输出路径（默认覆盖原文件）")
    
    # diff
    p_diff = sub.add_parser("diff", help="比较两个 DOCX 差异")
    p_diff.add_argument("file_a", help="第一个 DOCX 文件")
    p_diff.add_argument("file_b", help="第二个 DOCX 文件")
    
    # verify
    p_verify = sub.add_parser("verify", help="验证 DOCX 完整性")
    p_verify.add_argument("file", help="DOCX 文件路径")
    
    args = parser.parse_args()
    
    dispatch = {
        "analyze": cmd_analyze,
        "map": cmd_map,
        "replace": cmd_replace,
        "diff": cmd_diff,
        "verify": cmd_verify,
    }
    
    dispatch[args.command](args)


if __name__ == "__main__":
    main()
