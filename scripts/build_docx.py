#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown to DOCX builder with OMML equations, three-line tables with captions."""
import sys
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')
import argparse, re, os
from pathlib import Path

# Dependencies
for pkg, mod in [("python-docx","docx"),("latex2mathml","latex2mathml.converter"),("mathml2omml","mathml2omml"),("lxml","lxml")]:
    try: __import__(mod)
    except ImportError:
        print(f"[ERROR] Missing: {pkg}"); sys.exit(1)

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from latex2mathml.converter import convert as l2m_convert
from mathml2omml import convert as m2o
from lxml import etree

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# ── OMML Pipeline ────────────────────────────────────────────
def latex_to_omml(latex):
    latex = latex.strip()
    if not latex: return ""
    try: return m2o(l2m_convert(latex))
    except Exception as e:
        print(f"  [WARN] LaTeX->OMML: {latex[:50]}... ({e})")
        return ""

def inject_omml(para, omml_xml):
    if not omml_xml: return False
    try:
        omml_xml = omml_xml.replace("<m:oMath>", f'<m:oMath xmlns:m="{MATH_NS}">')
        para._p.append(etree.fromstring(omml_xml.encode("utf-8")))
        return True
    except Exception as e:
        print(f"  [WARN] OMML inject: {e}")
        return False

# ── Markdown cleanup ─────────────────────────────────────────
def clean_md_text(text):
    """Remove Markdown syntax that should not appear in the final DOCX text."""
    text = str(text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = text.replace("\\*", "*").replace("\\_", "_").replace("\\`", "`")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^\s*>\s*", "", text)
    text = re.sub(r"^\s*[-*+]\s+", "", text)
    text = re.sub(r"^\s*\d+[.)]\s+", "", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(?<!\w)(\*|_)([^*_]+)\1(?!\w)", r"\2", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def add_markdown_runs(paragraph, text, size=12):
    """Add cleaned text while preserving simple bold markdown."""
    text = str(text)
    pos = 0
    pattern = re.compile(r"(\*\*|__)(.+?)\1")
    for m in pattern.finditer(text):
        before = clean_md_text(text[pos:m.start()])
        if before:
            r = paragraph.add_run(before)
            set_font(r, size)
        bold_text = clean_md_text(m.group(2))
        if bold_text:
            r = paragraph.add_run(bold_text)
            set_font(r, size, bold=True)
        pos = m.end()
    rest = clean_md_text(text[pos:])
    if rest:
        r = paragraph.add_run(rest)
        set_font(r, size)

def count_substantive_chars(text):
    """Approximate paper substance, excluding Markdown/code scaffolding."""
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"\$\$.*?\$\$", "", text, flags=re.DOTALL)
    text = re.sub(r"\|[-:\s|]+\|?", "", text)
    text = clean_md_text(text)
    return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))

# ── Code Appendix Handling ───────────────────────────────────
def add_code_block(doc, code_text, language='python'):
    """
    添加代码块到 DOCX（等宽字体、保留缩进、浅灰背景）
    
    Args:
        doc: Document对象
        code_text: 代码文本
        language: 编程语言（用于标注，暂不影响格式）
    """
    # 添加代码段落
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(9)
    
    # 设置浅灰色背景
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    run._element.rPr.append(shading)

def collect_code_appendix(outputs_dir):
    """
    从 outputs/ 收集所有代码文件，按子问题分组
    
    Returns:
        dict: {subquestion: [file_paths]}
    """
    outputs_path = Path(outputs_dir)
    code_files = {}
    
    # 扫描 outputs/ 目录下的 Q* 子目录
    for q_dir in sorted(outputs_path.glob("Q*")):
        if not q_dir.is_dir():
            continue
        
        subq_name = q_dir.name
        subq_files = []
        
        # 收集该子问题下的所有 .py 文件
        for py_file in sorted(q_dir.glob("*.py")):
            if py_file.name.startswith("__"):
                continue  # 跳过 __init__.py 等
            subq_files.append(py_file)
        
        if subq_files:
            code_files[subq_name] = subq_files
    
    # 如果没有 Q* 目录，尝试直接在 outputs/ 下查找代码
    if not code_files:
        for py_file in sorted(outputs_path.glob("*.py")):
            if py_file.name.startswith("__"):
                continue
            code_files.setdefault("Main", []).append(py_file)
    
    return code_files

def add_code_appendix(doc, outputs_dir, is_english=False):
    """
    添加代码附录章节到论文末尾
    
    Args:
        doc: Document对象
        outputs_dir: outputs目录路径
        is_english: 是否英文论文
    """
    code_files = collect_code_appendix(outputs_dir)
    
    if not code_files:
        print("  [INFO] No code files found for appendix")
        return
    
    # 添加附录主标题
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    title_text = "Appendices" if is_english else "附录"
    r = p.add_run(title_text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    r.font.bold = True
    
    # 为每个子问题添加附录
    appendix_labels = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for idx, (subq, files) in enumerate(code_files.items()):
        if idx >= len(appendix_labels):
            break
        
        label = appendix_labels[idx]
        
        # 附录子标题
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        if is_english:
            heading_text = f"Appendix {label}: {subq} Code"
        else:
            heading_text = f"附录{label}: {subq}代码"
        r = p.add_run(heading_text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(14)
        r.font.bold = True
        
        # 添加每个文件的代码
        for file_path in files:
            # 文件名标题
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            r = p.add_run(f"📄 {file_path.name}")
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(11)
            r.font.bold = True
            
            # 读取并添加代码
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    code_content = f.read()
                
                # 限制单个文件最大长度（避免论文过长）
                max_lines = 500
                lines = code_content.split('\n')
                if len(lines) > max_lines:
                    code_content = '\n'.join(lines[:max_lines])
                    code_content += f"\n\n# ... (省略剩余 {len(lines)-max_lines} 行)"
                
                add_code_block(doc, code_content)
                print(f"  [CODE] Added {file_path.name} ({len(lines)} lines)")
            
            except Exception as e:
                print(f"  [WARN] Failed to read {file_path}: {e}")
                p = doc.add_paragraph()
                r = p.add_run(f"[ERROR] Failed to read file: {e}")
                r.font.size = Pt(9)
                r.font.color.rgb = (255, 0, 0)

# ── Three-line Table ─────────────────────────────────────────
def _add_table_cell_text(cell, text):
    """Add text to a table cell, converting $...$ inline formulas to OMML."""
    INLINE_EQ = re.compile(r"\$(.+?)\$")
    segments = INLINE_EQ.split(str(text))
    
    # Clear cell
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for k, seg in enumerate(segments):
        if k % 2 == 1:
            # Equation segment - convert to OMML
            omml = latex_to_omml(seg)
            if omml:
                inject_omml(p, omml)
            else:
                r = p.add_run("$" + clean_md_text(seg) + "$")
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(9)
        else:
            # Text segment
            cleaned = clean_md_text(seg)
            if cleaned:
                r = p.add_run(cleaned)
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(9)

def add_three_line_table(doc, headers, rows, caption_text=None):
    """CUMCM three-line table with optional caption."""
    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(clean_md_text(caption_text))
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10)
        r.font.bold = True

    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    # Borders: top+bottom thick, no vertical
    tbl = table._tbl
    tblPr = tbl.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr")
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for tag, val, sz in [("top","single","12"),("bottom","single","12"),
                          ("left","none","0"),("right","none","0"),
                          ("insideH","none","0"),("insideV","none","0")]:
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:val"),val); e.set(qn("w:sz"),sz)
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),"000000"); borders.append(e)
    tblPr.append(borders)

    # Header row - with inline formula support
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _add_table_cell_text(cell, h)
        # Make header bold
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
        # Thin bottom border on header cells
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        be = OxmlElement("w:bottom"); be.set(qn("w:val"),"single")
        be.set(qn("w:sz"),"4"); be.set(qn("w:color"),"000000")
        tcBorders.append(be); tcPr.append(tcBorders)

    # Data rows - with inline formula support
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            _add_table_cell_text(cell, str(val))

    doc.add_paragraph()  # spacing after table
    return table

# ── Markdown Table Parser ────────────────────────────────────
def parse_md_table(lines, idx):
    """Parse Markdown table. Returns (headers, rows, next_idx) or (None,None,idx)."""
    if idx >= len(lines): return None, None, idx
    hl = lines[idx].strip()
    if "|" not in hl: return None, None, idx
    headers = [h.strip() for h in hl.split("|") if h.strip()]
    if idx+1 >= len(lines): return None, None, idx
    sep = lines[idx+1].strip()
    if not re.match(r"^[\|\s\-:]+$", sep): return None, None, idx

    rows = []
    j = idx + 2
    while j < len(lines):
        rl = lines[j].strip()
        if "|" not in rl: break
        vals = [v.strip() for v in rl.split("|") if v.strip()]
        if len(vals) >= len(headers):
            rows.append(vals[:len(headers)])
        elif len(vals) > 0:
            rows.append(vals + [""]*(len(headers)-len(vals)))
        j += 1
    return headers, rows, j

# ── Competition Profiles ─────────────────────────────────────
COMPETITION_PROFILES = {
    "51mcm": {
        "cn_body": "宋体", "cn_heading": "黑体", "en_body": "Times New Roman",
        "margin_cm": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
        "min_chars": 9000, "language": "zh",
    },
    "cumcm": {
        "cn_body": "宋体", "cn_heading": "黑体", "en_body": "Times New Roman",
        "margin_cm": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
        "min_chars": 9000, "language": "zh",
    },
    "mcm-icm": {
        "cn_body": None, "cn_heading": None, "en_body": "Times New Roman",
        "margin_cm": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
        "min_chars": 0, "language": "en",
    },
}

# ── Font helper ──────────────────────────────────────────────
def set_font(run, size=12, bold=False, cn="宋体", en="Times New Roman"):
    run.font.name = en
    if cn:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size); run.font.bold = bold

# ── Main Builder ─────────────────────────────────────────────
def build_paper(md_path, output_path, competition="cumcm"):
    profile = COMPETITION_PROFILES.get(competition, COMPETITION_PROFILES["cumcm"])
    cn_body = profile["cn_body"] or "Times New Roman"
    cn_heading = profile["cn_heading"] or "Times New Roman"
    en_body = profile["en_body"]
    margins = profile["margin_cm"]
    min_chars = profile["min_chars"]
    is_english = profile["language"] == "en"

    with open(md_path, "r", encoding="utf-8") as f:
        raw = f.read()
    lines = raw.split("\n")
    source_chars = count_substantive_chars(raw)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(margins["top"]); s.bottom_margin = Cm(margins["bottom"])
        s.left_margin = Cm(margins["left"]); s.right_margin = Cm(margins["right"])

    i = 0; eq_ok = eq_fail = tbl_n = fig_n = 0
    INLINE_EQ = re.compile(r"\$(.+?)\$")
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Code fence - skip fenced code content entirely.
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # Skip empty
        if not line: i += 1; continue

        # Horizontal rule / separator - skip silently
        if line.strip() in ("---", "***", "___"):
            i += 1; continue

        # Heading
        if line.startswith("## "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
            r = p.add_run(clean_md_text(line[3:])); set_font(r, 14, True, cn_heading, en_body); i += 1; continue
        if line.startswith("### "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
            r = p.add_run(clean_md_text(line[4:])); set_font(r, 12, True, cn_heading, en_body); i += 1; continue
        if line.startswith("#### "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
            r = p.add_run(clean_md_text(line[5:])); set_font(r, 12, True, cn_body, en_body); i += 1; continue
        if line.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(40)
            r = p.add_run(clean_md_text(line[2:])); set_font(r, 18, True, cn_heading, en_body); i += 1; continue

        # Table
        if "|" in line and i+1 < len(lines) and re.match(r"^[\|\s\-:]+$", lines[i+1].strip()):
            caption = None
            j = i - 1
            while j >= 0 and not lines[j].strip(): j -= 1
            if j >= 0:
                prev = lines[j].strip()
                cm = re.match(r"(?:\*\*)?表\s*(\d+)\s*[：:\s]?\s*(.+?)(?:\*\*)?$", prev)
                en_cm = re.match(r"(?:\*\*)?Table\s+(\d+)\s*[.:：\s]\s*(.+?)(?:\*\*)?$", prev, re.IGNORECASE)
                if cm:
                    if is_english:
                        caption = f"Table {cm.group(1)}: {clean_md_text(cm.group(2))}"
                    else:
                        caption = f"表{cm.group(1)} {clean_md_text(cm.group(2))}"
                    lines[j] = ""
                elif en_cm:
                    caption = f"Table {en_cm.group(1)}: {clean_md_text(en_cm.group(2))}"
                    lines[j] = ""

            headers, rows, next_i = parse_md_table(lines, i)
            if headers and rows:
                tbl_n += 1
                if caption is None:
                    caption = f"Table {tbl_n}" if is_english else f"表{tbl_n}"
                add_three_line_table(doc, headers, rows, caption)
                i = next_i; continue

        # Display equation: $$...$$ (single or multi-line)
        if line.startswith("$$"):
            if len(line) > 4 and line.endswith("$$"):
                latex = line[2:-2].strip()
            else:
                latex_lines = []
                j = i + 1
                while j < len(lines) and not lines[j].strip() == "$$":
                    latex_lines.append(lines[j])
                    j += 1
                latex = "\n".join(latex_lines).strip()
                i = j
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            omml = latex_to_omml(latex)
            if omml and inject_omml(p, omml):
                eq_ok += 1; print(f"  [EQ-{eq_ok}] {latex[:60]}...")
            else:
                eq_fail += 1
                placeholder = f"[Eq: {latex[:40]}...]" if is_english else f"[公式: {latex[:40]}...]"
                r = p.add_run(placeholder); set_font(r, 10)
            i += 1; continue

        # FIGURE: [FIGURE: path | caption]
        fm = re.match(r"\[FIGURE:\s*(.+?)\s*\|\s*(.+?)\]", line)
        if fm:
            fname = fm.group(1).strip(); cap_text = clean_md_text(fm.group(2))
            if os.path.exists(fname):
                fig_n += 1
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    r = p.add_run(); r.add_picture(fname, width=Inches(5.0))
                    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(cap_text); set_font(rc, 10, cn=cn_body)
                    print(f"  [FIG-{fig_n}] {os.path.basename(fname)}")
                except Exception as e:
                    print(f"  [WARN] Image failed: {fname} - {e}")
            else:
                print(f"  [WARN] Image not found: {fname}")
            i += 1; continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24) if not is_english else Pt(0)
        p.paragraph_format.line_spacing = 1.5

        eq_segments = INLINE_EQ.split(line)
        for k, seg in enumerate(eq_segments):
            if k % 2 == 1:
                omml = latex_to_omml(seg)
                if omml:
                    inject_omml(p, omml); eq_ok += 1
                else:
                    r = p.add_run(f"${seg}$"); set_font(r, 10)
            else:
                add_markdown_runs(p, seg, 12)
        i += 1

    # Content length check
    all_text = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    chars = len(all_text)
    print(f"\n[CHECK] Total chars: {chars} | source substance: {source_chars}")
    print(f"        competition: {competition} | language: {profile['language']}")
    if min_chars > 0:
        print(f"        hard minimum {min_chars}; no artificial upper/target length")
        if chars < min_chars:
            print(f"  HARD FAIL: {chars} < {min_chars}. The draft is too thin to submit.")
            print("  Expand with derivations, baseline comparison, robustness, figure interpretation, and error analysis.")
            return None
        print("  PASS: content clears the minimum. Do not pad; let the problem depth determine final length.")
    else:
        print("  No minimum char requirement for this competition type.")
    
    # 添加代码附录（在保存前）
    print("\n[Appendix] Collecting code files...")
    outputs_dir = Path(md_path).parent.parent / "outputs"
    if outputs_dir.exists():
        add_code_appendix(doc, outputs_dir, is_english)
    else:
        print("  [WARN] outputs/ directory not found, skipping code appendix")

    doc.save(output_path)
    pages = chars / 900
    print(f"[DONE] {output_path}")
    print(f"  Equations: {eq_ok} OK, {eq_fail} FAIL")
    print(f"  Tables: {tbl_n} | Figures: {fig_n}")
    print(f"  Total chars: {chars} | Pages: ~{pages:.0f}")
    return output_path

if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Markdown to DOCX builder with competition-aware formatting")
    ap.add_argument("markdown")
    ap.add_argument("output")
    ap.add_argument("--competition", "-c", default="cumcm",
                    choices=list(COMPETITION_PROFILES.keys()),
                    help="Competition type: cumcm (default), 51mcm, mcm-icm")
    args = ap.parse_args()
    build_paper(args.markdown, args.output, competition=args.competition)


