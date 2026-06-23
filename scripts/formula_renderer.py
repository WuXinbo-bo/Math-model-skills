#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
formula_renderer.py — 独立公式渲染引擎
LaTeX → MathML → OMML 转换，支持单条/批量/DOCX注入

功能:
  1. latex_to_omml()     — LaTeX 字符串 → OMML XML
  2. render_to_docx()    — 单条公式 → DOCX 文件
  3. batch_render()      — 批量渲染（从文件读取公式列表）
  4. inject_into_docx()  — 将 OMML 公式注入已有 DOCX 段落
  5. validate()          — 验证 LaTeX 公式可转换性

依赖: latex2mathml, mathml2omml, lxml, python-docx
"""
import sys
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

import argparse
import json
import re
import os
from pathlib import Path

# ── 依赖检查 ──────────────────────────────────────────────────
_MISSING = []
for pkg, mod in [
    ("python-docx", "docx"),
    ("latex2mathml", "latex2mathml.converter"),
    ("mathml2omml", "mathml2omml"),
    ("lxml", "lxml"),
]:
    try:
        __import__(mod)
    except ImportError:
        _MISSING.append(pkg)

if _MISSING:
    print(f"[ERROR] 缺少依赖: {', '.join(_MISSING)}")
    print(f"        pip install {' '.join(_MISSING)}")
    sys.exit(1)

from latex2mathml.converter import convert as l2m_convert
from mathml2omml import convert as m2o
from lxml import etree
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

__version__ = "1.0.0"


# ═══════════════════════════════════════════════════════════════
#  核心转换
# ═══════════════════════════════════════════════════════════════

def latex_to_omml(latex: str) -> str:
    """LaTeX → OMML XML 字符串。失败返回空字符串。"""
    latex = latex.strip()
    if not latex:
        return ""
    try:
        mathml = l2m_convert(latex)
        return m2o(mathml)
    except Exception as e:
        print(f"  [WARN] LaTeX→OMML 失败: {latex[:60]}... ({e})")
        return ""


def latex_to_mathml(latex: str) -> str:
    """LaTeX → MathML XML 字符串。失败返回空字符串。"""
    latex = latex.strip()
    if not latex:
        return ""
    try:
        return l2m_convert(latex)
    except Exception as e:
        print(f"  [WARN] LaTeX→MathML 失败: {latex[:60]}... ({e})")
        return ""


# ═══════════════════════════════════════════════════════════════
#  DOCX 注入
# ═══════════════════════════════════════════════════════════════

def inject_omml_into_paragraph(paragraph, omml_xml: str) -> bool:
    """将 OMML XML 注入 python-docx 段落。成功返回 True。"""
    if not omml_xml:
        return False
    try:
        # 确保命名空间声明
        if 'xmlns:m=' not in omml_xml:
            omml_xml = omml_xml.replace(
                "<m:oMath>",
                f'<m:oMath xmlns:m="{MATH_NS}">'
            )
        omml_element = etree.fromstring(omml_xml.encode("utf-8"))
        paragraph._p.append(omml_element)
        return True
    except Exception as e:
        print(f"  [WARN] OMML 注入失败: {e}")
        return False


def render_to_docx(latex: str, output_path: str, font_size: int = 12) -> bool:
    """将单条 LaTeX 公式渲染为独立 DOCX 文件。"""
    omml = latex_to_omml(latex)
    if not omml:
        print(f"[FAIL] 无法转换公式: {latex[:80]}")
        return False

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if inject_omml_into_paragraph(p, omml):
        doc.save(output_path)
        print(f"[OK] 公式已保存至 {output_path}")
        return True
    else:
        print(f"[FAIL] OMML 注入失败")
        return False


def render_batch_to_docx(formulas: list, output_path: str, font_size: int = 12) -> dict:
    """批量渲染公式到一个 DOCX 文件，每条公式一个段落。

    Args:
        formulas: LaTeX 公式列表
        output_path: 输出 DOCX 路径
        font_size: 字号（预留，OMML 自身控制大小）

    Returns:
        {"ok": int, "fail": int, "total": int}
    """
    doc = Document()
    ok = fail = 0

    for idx, latex in enumerate(formulas, 1):
        latex = latex.strip()
        if not latex:
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        omml = latex_to_omml(latex)
        if omml and inject_omml_into_paragraph(p, omml):
            ok += 1
            print(f"  [{ok}] {latex[:60]}...")
        else:
            fail += 1
            # 回退为纯文本占位符
            r = p.add_run(f"[公式渲染失败: {latex[:40]}...]")
            r.font.size = Pt(10)
            r.font.color.rgb = None  # 默认颜色

    doc.save(output_path)
    print(f"\n[DONE] {output_path}")
    print(f"  成功: {ok} | 失败: {fail} | 总计: {ok + fail}")
    return {"ok": ok, "fail": fail, "total": ok + fail}


# ═══════════════════════════════════════════════════════════════
#  公式提取
# ═══════════════════════════════════════════════════════════════

def extract_formulas_from_text(text: str) -> list:
    """从文本中提取所有 LaTeX 公式（$...$ 和 $$...$$）。

    Returns:
        [(latex_str, is_display), ...]
        is_display: True 表示块级公式（$$...$$），False 表示行内（$...$）
    """
    results = []

    # 块级公式 $$...$$
    for m in re.finditer(r'\$\$(.+?)\$\$', text, re.DOTALL):
        results.append((m.group(1).strip(), True))

    # 行内公式 $...$ （排除已被 $$ 匹配的）
    # 先移除 $$...$$ 再匹配 $...$
    text_no_display = re.sub(r'\$\$.+?\$\$', '', text, flags=re.DOTALL)
    for m in re.finditer(r'\$(.+?)\$', text_no_display):
        results.append((m.group(1).strip(), False))

    return results


def extract_formulas_from_file(filepath: str) -> list:
    """从文件中提取所有 LaTeX 公式。"""
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    return extract_formulas_from_text(text)


# ═══════════════════════════════════════════════════════════════
#  验证
# ═══════════════════════════════════════════════════════════════

def validate_formula(latex: str) -> dict:
    """验证 LaTeX 公式是否可成功转换为 OMML。

    Returns:
        {"valid": bool, "latex": str, "mathml": str, "omml": str, "error": str|None}
    """
    result = {
        "valid": False,
        "latex": latex.strip(),
        "mathml": "",
        "omml": "",
        "error": None,
    }

    try:
        mathml = l2m_convert(latex.strip())
        result["mathml"] = mathml
    except Exception as e:
        result["error"] = f"LaTeX→MathML: {e}"
        return result

    try:
        omml = m2o(mathml)
        result["omml"] = omml
        result["valid"] = True
    except Exception as e:
        result["error"] = f"MathML→OMML: {e}"

    return result


def validate_file(filepath: str) -> dict:
    """验证文件中所有公式的可转换性。"""
    formulas = extract_formulas_from_file(filepath)
    results = []
    ok = fail = 0

    for latex, is_display in formulas:
        v = validate_formula(latex)
        v["display"] = is_display
        results.append(v)
        if v["valid"]:
            ok += 1
        else:
            fail += 1
            print(f"  [FAIL] {latex[:50]}... — {v['error']}")

    return {
        "file": filepath,
        "total": len(results),
        "ok": ok,
        "fail": fail,
        "details": results,
    }


# ═══════════════════════════════════════════════════════════════
#  CLI 命令
# ═══════════════════════════════════════════════════════════════

def cmd_render(args):
    """渲染单条公式"""
    latex = args.latex
    fmt = args.format

    if fmt == "omml":
        result = latex_to_omml(latex)
        if result:
            print(result)
        else:
            print("[FAIL] 转换失败", file=sys.stderr)
            sys.exit(1)
    elif fmt == "mathml":
        result = latex_to_mathml(latex)
        if result:
            print(result)
        else:
            print("[FAIL] 转换失败", file=sys.stderr)
            sys.exit(1)
    elif fmt == "docx":
        output = args.output or "formula_output.docx"
        if not render_to_docx(latex, output):
            sys.exit(1)


def cmd_batch(args):
    """批量渲染公式"""
    input_path = args.input
    output = args.output or "formulas_batch.docx"

    # 读取公式列表
    if input_path.endswith(".json"):
        with open(input_path, "r", encoding="utf-8") as f:
            formulas = json.load(f)
    else:
        # 每行一条公式
        with open(input_path, "r", encoding="utf-8") as f:
            formulas = [line.strip() for line in f if line.strip() and not line.startswith("#")]

    result = render_batch_to_docx(formulas, output)
    if result["fail"] > 0 and args.strict:
        sys.exit(1)


def cmd_extract(args):
    """从文件中提取公式"""
    filepath = args.input
    output = args.output

    formulas = extract_formulas_from_file(filepath)
    print(f"[INFO] 从 {filepath} 提取到 {len(formulas)} 条公式")

    if output:
        if output.endswith(".json"):
            data = [{"latex": f, "display": d} for f, d in formulas]
            with open(output, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        else:
            with open(output, "w", encoding="utf-8") as f:
                for latex, is_display in formulas:
                    prefix = "$$" if is_display else "$"
                    f.write(f"{prefix}{latex}{prefix}\n")
        print(f"[OK] 已保存至 {output}")
    else:
        for idx, (latex, is_display) in enumerate(formulas, 1):
            tag = "DISPLAY" if is_display else "INLINE"
            print(f"  [{idx}] ({tag}) {latex}")


def cmd_validate(args):
    """验证公式可转换性"""
    if args.input:
        # 验证文件中的所有公式
        result = validate_file(args.input)
        print(f"\n[RESULT] {result['file']}")
        print(f"  总计: {result['total']} | 通过: {result['ok']} | 失败: {result['fail']}")
        if result["fail"] > 0:
            sys.exit(1)
    else:
        # 验证单条公式
        v = validate_formula(args.latex)
        if v["valid"]:
            print(f"[OK] 公式可转换")
            if args.verbose:
                print(f"  MathML: {v['mathml'][:200]}...")
                print(f"  OMML:   {v['omml'][:200]}...")
        else:
            print(f"[FAIL] {v['error']}")
            sys.exit(1)


# ═══════════════════════════════════════════════════════════════
#  CLI 入口
# ═══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="formula_renderer.py — 独立公式渲染引擎 (LaTeX→OMML)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 渲染单条公式为 OMML
  python formula_renderer.py render "E = mc^2"

  # 渲染为 MathML
  python formula_renderer.py render --format mathml "\\frac{a}{b}"

  # 渲染为 DOCX
  python formula_renderer.py render --format docx --output eq.docx "\\sum_{i=1}^{n} x_i"

  # 批量渲染（从文件）
  python formula_renderer.py batch formulas.txt --output batch.docx

  # 从 Markdown 提取公式
  python formula_renderer.py extract paper.md --output formulas.json

  # 验证公式
  python formula_renderer.py validate "\\int_0^1 f(x) dx"
  python formula_renderer.py validate --input paper.md
        """,
    )
    sub = parser.add_subparsers(dest="command", help="子命令")

    # ── render ──
    p_render = sub.add_parser("render", help="渲染单条 LaTeX 公式")
    p_render.add_argument("latex", help="LaTeX 公式字符串")
    p_render.add_argument("--format", "-f", choices=["omml", "mathml", "docx"],
                          default="omml", help="输出格式 (默认: omml)")
    p_render.add_argument("--output", "-o", help="输出文件路径 (仅 docx 格式)")
    p_render.set_defaults(func=cmd_render)

    # ── batch ──
    p_batch = sub.add_parser("batch", help="批量渲染公式到 DOCX")
    p_batch.add_argument("input", help="输入文件 (.txt 每行一条 / .json 数组)")
    p_batch.add_argument("--output", "-o", help="输出 DOCX 路径")
    p_batch.add_argument("--strict", action="store_true",
                         help="有任何失败则退出码非零")
    p_batch.set_defaults(func=cmd_batch)

    # ── extract ──
    p_extract = sub.add_parser("extract", help="从文件中提取 LaTeX 公式")
    p_extract.add_argument("input", help="输入文件路径 (Markdown/LaTeX/纯文本)")
    p_extract.add_argument("--output", "-o",
                           help="输出文件 (.json 或 .txt，不指定则打印)")
    p_extract.set_defaults(func=cmd_extract)

    # ── validate ──
    p_validate = sub.add_parser("validate", help="验证公式可转换性")
    p_validate.add_argument("latex", nargs="?", default="",
                            help="LaTeX 公式 (不指定则用 --input)")
    p_validate.add_argument("--input", "-i", help="验证文件中所有公式")
    p_validate.add_argument("--verbose", "-v", action="store_true",
                            help="显示中间转换结果")
    p_validate.set_defaults(func=cmd_validate)

    args = parser.parse_args()
    if not args.command:
        parser.print_help()
        sys.exit(0)

    args.func(args)


if __name__ == "__main__":
    main()
